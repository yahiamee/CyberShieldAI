from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Read the English document
input_file = 'CyberShield_AI_Project_Report.docx'
output_file = 'CyberShield_AI_Project_Report_Arabic.docx'

doc = Document(input_file)

# Translation dictionary for common terms
translations = {
    'CyberShield AI': 'CyberShield AI',
    'Project Report': 'تقرير المشروع',
    'Intelligent Web Security Scanner Using Artificial Intelligence': 'نظام ذكي لفحص أمان المواقع الإلكترونية باستخدام الذكاء الاصطناعي',
    'Introduction': 'مقدمة',
    'Project Overview': 'نظرة عامة على المشروع',
    'Week 1: Planning and Analysis': 'الأسبوع الأول: التخطيط والتحليل',
    'Week 2: Design and Architecture': 'الأسبوع الثاني: التصميم والهندسة المعمارية',
    'Week 3: Development Environment Setup and Python Installation': 'الأسبوع الثالث: تهيئة البيئة التطويرية وتثبيت Python',
    'Week 4: Basic System Development and Flask': 'الأسبوع الرابع: تطوير النظام الأساسي و Flask',
    'Week 5: Basic Scanning System Development': 'الأسبوع الخامس: تطوير نظام الفحص الأساسي',
    'Week 6: PDF Report Generation System Development': 'الأسبوع السادس: تطوير نظام توليد التقارير PDF',
    'Week 7: OWASP ZAP Integration': 'الأسبوع السابع: تكامل OWASP ZAP',
    'Week 8: Automatic ZAP Management': 'الأسبوع الثامن: إدارة ZAP التلقائية',
    'Week 9: Interface Improvements and Error Handling': 'الأسبوع التاسع: تحسين الواجهات ومعالجة الأخطاء',
    'Week 10: Report Improvements': 'الأسبوع العاشر: تحسين التقارير',
    'Week 11: OpenAI Integration': 'الأسبوع الحادي عشر: تكامل OpenAI',
    'Week 12: Multi-language Support': 'الأسبوع الثاني عشر: دعم اللغات المتعددة',
    'Week 13: Adding Vulnerabilities Section to Home Page': 'الأسبوع الثالث عشر: إضافة قسم الثغرات في الصفحة الرئيسية',
    'Week 14: Final Testing and Documentation': 'الأسبوع الرابع عشر: الاختبار النهائي والتوثيق',
    'Final Project Structure': 'البنية النهائية للمشروع',
    'Libraries Used': 'المكتبات المستخدمة',
    'Final System Features': 'المميزات النهائية للنظام',
    'System Operation Steps': 'الخطوات المتبعة لتشغيل النظام',
    'Conclusion': 'الخلاصة',
    'Table of Contents': 'جدول المحتويات',
}

def translate_text(text):
    """Translate English text to Arabic"""
    if not text or not text.strip():
        return text
    
    # Check if it's an image placeholder
    if '[Image Placeholder:' in text:
        # Translate the description part
        if 'Shows' in text:
            text = text.replace('Shows', 'يوضح')
        if 'Project Analysis Diagram' in text:
            text = text.replace('Project Analysis Diagram', 'مخطط تحليل المشروع')
        if 'System Architecture Diagram' in text:
            text = text.replace('System Architecture Diagram', 'مخطط البنية المعمارية للنظام')
        if 'Database ERD Diagram' in text:
            text = text.replace('Database ERD Diagram', 'مخطط قاعدة البيانات ERD')
        if 'User Interface Designs' in text:
            text = text.replace('User Interface Designs', 'تصميمات واجهة المستخدم')
        if 'Python Installation Screenshot' in text:
            text = text.replace('Python Installation Screenshot', 'لقطة شاشة لتثبيت Python')
        if 'Library Installation Screenshot' in text:
            text = text.replace('Library Installation Screenshot', 'لقطة شاشة لتثبيت المكتبات')
        if 'Project Folder Structure' in text:
            text = text.replace('Project Folder Structure', 'هيكل مجلدات المشروع')
        if 'Home Page Screenshot' in text:
            text = text.replace('Home Page Screenshot', 'لقطة شاشة للصفحة الرئيسية')
        if 'Login Page Screenshot' in text:
            text = text.replace('Login Page Screenshot', 'لقطة شاشة لصفحة تسجيل الدخول')
        if 'Scan Page Screenshot' in text:
            text = text.replace('Scan Page Screenshot', 'لقطة شاشة لصفحة الفحص')
        if 'Results Page Screenshot' in text:
            text = text.replace('Results Page Screenshot', 'لقطة شاشة لصفحة النتائج')
        if 'PDF Report Screenshot' in text:
            text = text.replace('PDF Report Screenshot', 'لقطة شاشة لتقرير PDF')
        if 'Complete PDF Report Example' in text:
            text = text.replace('Complete PDF Report Example', 'مثال على تقرير PDF مكتمل')
        if 'ZAP Settings Screenshot' in text:
            text = text.replace('ZAP Settings Screenshot', 'لقطة شاشة لإعدادات ZAP')
        if 'ZAP Results Screenshot' in text:
            text = text.replace('ZAP Results Screenshot', 'لقطة شاشة لنتائج ZAP')
        if 'OpenAI Settings Screenshot' in text:
            text = text.replace('OpenAI Settings Screenshot', 'لقطة شاشة لإعدادات OpenAI')
        if 'OpenAI Analysis Screenshot' in text:
            text = text.replace('OpenAI Analysis Screenshot', 'لقطة شاشة لتحليل OpenAI')
        if 'System in Arabic Screenshot' in text:
            text = text.replace('System in Arabic Screenshot', 'لقطة شاشة للنظام باللغة العربية')
        if 'Language Switcher Screenshot' in text:
            text = text.replace('Language Switcher Screenshot', 'لقطة شاشة لمبدل اللغة')
        if 'System in English Screenshot' in text:
            text = text.replace('System in English Screenshot', 'لقطة شاشة للنظام باللغة الإنجليزية')
        if 'Vulnerabilities Section Screenshot' in text:
            text = text.replace('Vulnerabilities Section Screenshot', 'لقطة شاشة لقسم الثغرات')
        if 'System Testing Screenshot' in text:
            text = text.replace('System Testing Screenshot', 'لقطة شاشة لاختبار النظام')
        if 'Documentation Screenshot' in text:
            text = text.replace('Documentation Screenshot', 'لقطة شاشة للوثائق')
        if 'Complete Project Structure' in text:
            text = text.replace('Complete Project Structure', 'هيكل المشروع الكامل')
        if 'Database Diagram' in text:
            text = text.replace('Database Diagram', 'مخطط قاعدة البيانات')
        if 'Libraries List' in text:
            text = text.replace('Libraries List', 'قائمة المكتبات')
        if 'Main Features Screenshot' in text:
            text = text.replace('Main Features Screenshot', 'لقطة شاشة للميزات الرئيسية')
        if 'Interface Screenshot' in text:
            text = text.replace('Interface Screenshot', 'لقطة شاشة للواجهة')
        if 'System Running Screenshot' in text:
            text = text.replace('System Running Screenshot', 'لقطة شاشة لتشغيل النظام')
        if 'Final System Screenshot' in text:
            text = text.replace('Final System Screenshot', 'لقطة شاشة نهائية للنظام')
        if 'Final Report Example' in text:
            text = text.replace('Final Report Example', 'مثال على تقرير نهائي')
        return text
    
    # Direct translations
    for eng, arb in translations.items():
        if eng in text:
            text = text.replace(eng, arb)
    
    return text

# Create new document
new_doc = Document()

# Copy document settings
new_doc.core_properties.title = doc.core_properties.title
new_doc.core_properties.author = doc.core_properties.author

# Copy sections settings
for i, section in enumerate(doc.sections):
    if i < len(new_doc.sections):
        new_section = new_doc.sections[i]
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin

# Process all paragraphs
for para in doc.paragraphs:
    # Check if paragraph has runs
    if para.runs:
        new_para = new_doc.add_paragraph()
        new_para.alignment = para.alignment
        new_para.style = para.style
        
        for run in para.runs:
            text = run.text
            translated_text = translate_text(text)
            
            new_run = new_para.add_run(translated_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
    else:
        # Empty paragraph
        new_para = new_doc.add_paragraph()
        new_para.alignment = para.alignment
        new_para.style = para.style

# Process tables
for table in doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.style = table.style
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.rows[i].cells[j]
            for para in cell.paragraphs:
                new_para = new_cell.add_paragraph()
                new_para.alignment = para.alignment
                for run in para.runs:
                    text = run.text
                    translated_text = translate_text(text)
                    new_run = new_para.add_run(translated_text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size

# Save the new document
new_doc.save(output_file)
print(f"تم إنشاء الملف العربي بنجاح: {output_file}")


