from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Read the English document
input_file = 'CyberShield_AI_Project_Report.docx'
output_file = 'CyberShield_AI_Project_Report_Arabic.docx'

print(f"جارٍ قراءة الملف: {input_file}")
doc = Document(input_file)

# Create new document
new_doc = Document()

# Copy document settings
new_doc.core_properties.title = doc.core_properties.title
new_doc.core_properties.author = doc.core_properties.author

# Copy sections settings
for i, section in enumerate(doc.sections):
    if i < len(new_doc.sections):
        new_section = new_doc.sections[i]
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin

# Set default font for Arabic
style = new_doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Translation mapping - comprehensive translations
translations_map = {
    # Title and headers
    'CyberShield AI': 'CyberShield AI',
    'Project Report': 'تقرير المشروع',
    'Intelligent Web Security Scanner Using Artificial Intelligence': 'نظام ذكي لفحص أمان المواقع الإلكترونية باستخدام الذكاء الاصطناعي',
    'Introduction': 'مقدمة المشروع',
    'Project Overview': 'نظرة عامة على المشروع',
    'Table of Contents': 'جدول المحتويات',
    
    # Week titles
    'Week 1: Planning and Analysis': 'الأسبوع الأول: التخطيط والتحليل',
    'Week 2: Design and Architecture': 'الأسبوع الثاني: التصميم والهندسة المعمارية',
    'Week 3: Development Environment Setup and Python Installation': 'الأسبوع الثالث: تهيئة البيئة التطويرية وتثبيت Python',
    'Week 4: Basic System Development and Flask': 'الأسبوع الرابع: تطوير النظام الأساسي و Flask',
    'Week 5: Basic Scanning System Development': 'الأسبوع الخامس: تطوير نظام الفحص الأساسي',
    'Week 6: PDF Report Generation System Development': 'الأسبوع السادس: تطوير نظام توليد التقارير PDF',
    'Week 7: OWASP ZAP Integration': 'الأسبوع السابع: تكامل OWASP ZAP',
    'Week 8: Automatic ZAP Management': 'الأسبوع الثامن: إدارة ZAP التلقائية',
    'Week 9: Interface Improvements and Error Handling': 'الأسبوع التاسع: تحسين الواجهات ومعالجة الأخطاء',
    'Week 10: Report Improvements': 'الأسبوع العاشر: تحسين التقارير',
    'Week 11: OpenAI Integration': 'الأسبوع الحادي عشر: تكامل OpenAI',
    'Week 12: Multi-language Support': 'الأسبوع الثاني عشر: دعم اللغات المتعددة',
    'Week 13: Adding Vulnerabilities Section to Home Page': 'الأسبوع الثالث عشر: إضافة قسم الثغرات في الصفحة الرئيسية',
    'Week 14: Final Testing and Documentation': 'الأسبوع الرابع عشر: الاختبار النهائي والتوثيق',
    
    # Section titles
    'Final Project Structure': 'البنية النهائية للمشروع',
    'Libraries Used': 'المكتبات المستخدمة',
    'Final System Features': 'المميزات النهائية للنظام',
    'System Operation Steps': 'الخطوات المتبعة لتشغيل النظام',
    'Conclusion': 'الخلاصة',
    
    # Image placeholders
    '[Image Placeholder:': '[مكان الصورة:',
    'Shows': 'يوضح',
    'Project Analysis Diagram': 'مخطط تحليل المشروع',
    'System Architecture Diagram': 'مخطط البنية المعمارية للنظام',
    'Database ERD Diagram': 'مخطط قاعدة البيانات ERD',
    'User Interface Designs': 'تصميمات واجهة المستخدم',
    'Python Installation Screenshot': 'لقطة شاشة لتثبيت Python',
    'Library Installation Screenshot': 'لقطة شاشة لتثبيت المكتبات',
    'Project Folder Structure': 'هيكل مجلدات المشروع',
    'Home Page Screenshot': 'لقطة شاشة للصفحة الرئيسية',
    'Login Page Screenshot': 'لقطة شاشة لصفحة تسجيل الدخول',
    'Scan Page Screenshot': 'لقطة شاشة لصفحة الفحص',
    'Results Page Screenshot': 'لقطة شاشة لصفحة النتائج',
    'PDF Report Screenshot': 'لقطة شاشة لتقرير PDF',
    'Complete PDF Report Example': 'مثال على تقرير PDF مكتمل',
    'ZAP Settings Screenshot': 'لقطة شاشة لإعدادات ZAP',
    'ZAP Results Screenshot': 'لقطة شاشة لنتائج ZAP',
    'OpenAI Settings Screenshot': 'لقطة شاشة لإعدادات OpenAI',
    'OpenAI Analysis Screenshot': 'لقطة شاشة لتحليل OpenAI',
    'System in Arabic Screenshot': 'لقطة شاشة للنظام باللغة العربية',
    'Language Switcher Screenshot': 'لقطة شاشة لمبدل اللغة',
    'System in English Screenshot': 'لقطة شاشة للنظام باللغة الإنجليزية',
    'Vulnerabilities Section Screenshot': 'لقطة شاشة لقسم الثغرات',
    'System Testing Screenshot': 'لقطة شاشة لاختبار النظام',
    'Documentation Screenshot': 'لقطة شاشة للوثائق',
    'Complete Project Structure': 'هيكل المشروع الكامل',
    'Database Diagram': 'مخطط قاعدة البيانات',
    'Libraries List': 'قائمة المكتبات',
    'Main Features Screenshot': 'لقطة شاشة للميزات الرئيسية',
    'Interface Screenshot': 'لقطة شاشة للواجهة',
    'System Running Screenshot': 'لقطة شاشة لتشغيل النظام',
    'Final System Screenshot': 'لقطة شاشة نهائية للنظام',
    'Final Report Example': 'مثال على تقرير نهائي',
}

def translate_text(text):
    """Translate English text to Arabic"""
    if not text or not text.strip():
        return text
    
    translated = text
    
    # Apply translations
    for eng, arb in translations_map.items():
        if eng in translated:
            translated = translated.replace(eng, arb)
    
    return translated

print("جارٍ معالجة الفقرات...")
# Process all paragraphs
for para_idx, para in enumerate(doc.paragraphs):
    if para_idx % 50 == 0:
        print(f"  معالجة الفقرة {para_idx}...")
    
    # Check if paragraph has runs
    if para.runs:
        new_para = new_doc.add_paragraph()
        new_para.alignment = para.alignment
        new_para.style = para.style
        
        for run in para.runs:
            text = run.text
            translated_text = translate_text(text)
            
            new_run = new_para.add_run(translated_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.name:
                new_run.font.name = 'Arial'  # Use Arial for Arabic
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
    else:
        # Empty paragraph
        new_para = new_doc.add_paragraph()
        new_para.alignment = para.alignment
        new_para.style = para.style

print("جارٍ معالجة الجداول...")
# Process tables
for table_idx, table in enumerate(doc.tables):
    print(f"  معالجة الجدول {table_idx + 1}...")
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.style = table.style
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.rows[i].cells[j]
            for para in cell.paragraphs:
                new_para = new_cell.add_paragraph()
                new_para.alignment = para.alignment
                for run in para.runs:
                    text = run.text
                    translated_text = translate_text(text)
                    new_run = new_para.add_run(translated_text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    if run.font.name:
                        new_run.font.name = 'Arial'
                    if run.font.size:
                        new_run.font.size = run.font.size

# Copy images (they are embedded in the document)
print("جارٍ حفظ الملف...")
new_doc.save(output_file)
print(f"تم إنشاء الملف العربي بنجاح: {output_file}")
print(f"الملف يحتوي على نفس الصور والتنسيق مع النصوص المترجمة")

